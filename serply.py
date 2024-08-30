import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Diccionario Filosófico del Pragmatismo", page_icon="📚", layout="wide")

# Function to create the information column
def crear_columna_info():
    st.markdown("""
    ## Sobre esta aplicación

    Esta aplicación es un Diccionario Filosófico basado en el pensamiento de la Filosofía Pragmatista. Permite a los usuarios obtener definiciones de términos filosóficos según la interpretación de diversos autores de esta corriente.

    ### Cómo usar la aplicación:

    1. Elija un término filosófico de la lista predefinida o proponga su propio término.
    2. Haga clic en "Obtener definición" para generar las definiciones.
    3. Lea las definiciones y fuentes proporcionadas.
    4. Si lo desea, descargue un documento DOCX con toda la información.

    ### Autor y actualización:
    **Moris Polanco**, 27 ag 2024

    ### Cómo citar esta aplicación (formato APA):
    Polanco, M. (2024). *Diccionario Filosófico del Pragmatismo* [Aplicación web]. https://pragmatismo.streamlit.app

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar definiciones basadas en información disponible en línea. Siempre verifique la información con fuentes académicas para un análisis más profundo.
    """)

# Titles and Main Column
st.title("Diccionario Filosófico del Pragmatismo")

col1, col2 = st.columns([1, 2])

with col1:
    crear_columna_info()

with col2:
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPLY_API_KEY = st.secrets["SERPLY_API_KEY"]

    # 101 philosophical terms related to Pragmatism
    terminos_filosoficos = sorted([
        "Acción", "Aprecio", "Certeza", "Cognición", "Comunicación", "Complejidad", "Continuidad", 
        "Conocimiento", "Contexto", "Convicción", "Cuerpo", "Desarrollo", "Diálogo", "Diferencia", 
        "Dinamismo", "Discusión", "Duda", "Eficacia", "Empiricismo", "Enero", "Epistemología", 
        "Escepticismo", "Experiencia", "Experimentación", "Fallibilismo", "Funcionalismo", "Habitus", 
        "Interacción", "Interdependencia", "Interpretación", "Intuición", "Investigación", "Justificación", 
        "Lenguaje", "Mediación", "Método científico", "Naturaleza", "Objeción", "Objetividad", "Observación", 
        "Operacionalismo", "Paradigma", "Pensamiento", "Percepción", "Pluralismo", "Práctica", "Pragmática", 
        "Pragmatismo", "Proceso", "Producto", "Psicología", "Quehacer", "Racionalidad", "Razón práctica", 
        "Realismo", "Relativismo", "Relación", "Representación", "Resiliencia", "Resolución", "Responsabilidad", 
        "Sensibilidad", "Significado", "Signo", "Simbiosis", "Síntesis", "Sujeto conocedor", "Suposición", 
        "Teoría", "Transformación", "Utopía", "Valor", "Verdad", "Vitalidad", "Voluntarismo"
    ])

    def buscar_informacion(query):
        url = "https://api.serply.io/v1/scholar"
        params = {
            "q": f"{query} Filosofía Pragmatista"
        }
        headers = {
            'X-Api-Key': SERPLY_API_KEY,
            'Content-Type': 'application/json'
        }
        response = requests.get(url, headers=headers, params=params)
        return response.json()

    def generar_definicion(termino, contexto):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Contexto: {contexto}\n\nTérmino: {termino}\n\nProporciona una definición del término filosófico '{termino}' según el pensamiento de la Filosofía Pragmatista. La definición debe ser concisa pero informativa, similar a una entrada de diccionario. Si es posible, incluye una referencia a una obra específica que trate este concepto.\n\nDefinición:",
            "max_tokens": 2048,
            "temperature": 0,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 0,
            "stop": ["Término:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def create_docx(terminos_definiciones):
        doc = Document()
        doc.add_heading('Diccionario Filosófico - Pragmatismo', 0)

        for termino, (definicion, fuentes) in terminos_definiciones.items():
            doc.add_heading('Término', level=1)
            doc.add_paragraph(termino)

            doc.add_heading('Definición', level=2)
            doc.add_paragraph(definicion)

            doc.add_heading('Fuentes', level=1)
            for fuente in fuentes:
                doc.add_paragraph(fuente, style='List Bullet')

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')

        return doc

    st.write("Elige un término filosófico de la lista, propón tu propio término, o genera todos los artículos:")

    opcion = st.radio("", ["Elegir de la lista", "Proponer mi propio término", "Generar todos los artículos en batch"])

    if opcion == "Elegir de la lista":
        termino = st.selectbox("Selecciona un término:", terminos_filosoficos)
    elif opcion == "Proponer mi propio término":
        termino = st.text_input("Ingresa tu propio término filosófico:")
    else:
        termino = None

    if st.button("Obtener definición" if opcion != "Generar todos los artículos en batch" else "Generar todos los artículos"):
        if opcion == "Generar todos los artículos en batch":
            with st.spinner("Generando todas las definiciones en batch..."):
                terminos_definiciones = {}
                for termino in terminos_filosoficos:
                    # Buscar información relevante
                    resultados_busqueda = buscar_informacion(termino)
                    contexto = "\n".join([item.get("snippet", "") for item in resultados_busqueda.get("organic", [])])
                    fuentes = [item.get("link", "") for item in resultados_busqueda.get("organic", [])]

                    # Generar definición
                    definicion = generar_definicion(termino, contexto)

                    terminos_definiciones[termino] = (definicion, fuentes)

                # Crear y descargar el documento con todas las definiciones
                doc = create_docx(terminos_definiciones)
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Descargar todas las definiciones en DOCX",
                    data=buffer,
                    file_name="Diccionario_Pragmatismo_Completo.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        elif termino:
            with st.spinner("Buscando información y generando definición..."):
                # Buscar información relevante
                resultados_busqueda = buscar_informacion(termino)
                contexto = "\n".join([item.get("snippet", "") for item in resultados_busqueda.get("organic", [])])
                fuentes = [item.get("link", "") for item in resultados_busqueda.get("organic", [])]

                # Generar definición
                definicion = generar_definicion(termino, contexto)

                # Mostrar la definición
                st.subheader(f"Definición para el término: {termino}")
                st.markdown(definicion)

                # Botón para descargar el documento
                doc = create_docx({termino: (definicion, fuentes)})
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.download_button(
                    label="Descargar definición en DOCX",
                    data=buffer,
                    file_name=f"Definicion_{termino.replace(' ', '_')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.warning("Por favor, selecciona un término.")

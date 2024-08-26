import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO

# Set page configuration
st.set_page_config(page_title="Diccionario Filosófico del Pragmatismo", page_icon="📚", layout="wide")

# Function to create the information column
def crear_columna_info():
    st.markdown("""
    ## Sobre esta aplicación

    Esta aplicación es un Diccionario Filosófico basado en el pensamiento de la Filosofía Pragmatista. Permite a los usuarios obtener definiciones de términos filosóficos según la interpretación de diversos autores de esta corriente.

    ### Cómo usar la aplicación:

    1. Elija un término filosófico de la lista predefinida o proponga su propio término.
    2. Seleccione uno o más autores pragmatistas.
    3. Haga clic en "Obtener definición" para generar las definiciones.
    4. Lea las definiciones y fuentes proporcionadas.
    5. Si lo desea, descargue un documento DOCX con toda la información.

    ### Autor y actualización:
    **Moris Polanco**, 26 ag 2024

    ### Cómo citar esta aplicación (formato APA):
    Polanco, M. (2024). *Diccionario Filosófico del Pragmatismo* [Aplicación web]. https://dicpragmatismo.streamlit.app

    ---
    **Nota:** Esta aplicación utiliza inteligencia artificial para generar definiciones basadas en información disponible en línea. Siempre verifique la información con fuentes académicas para un análisis más profundo.
    """)

# Titles and Main Column
st.title("Diccionario Filosófico del Pragmatismo")

col1, col2 = st.columns([1, 2])

with col1:
    crear_columna_info()

with col2:
    TOGETHER_API_KEY = st.secrets["TOGETHER_API_KEY"]
    SERPER_API_KEY = st.secrets["SERPER_API_KEY"]

    # 101 philosophical terms related to Pragmatism
    terminos_filosoficos = sorted([
        "Acción", "Aprecio", "Certeza", "Cognición", "Comunicación", "Complejidad", "Continuidad", 
        "Conocimiento", "Contexto", "Convicción", "Cuerpo", "Desarrollo", "Diálogo", "Diferencia", 
        "Dinamismo", "Discusión", "Duda", "Eficacia", "Empiricismo", "Enero", "Epistemología", 
        "Escepticismo", "Experiencia", "Experimentación", "Fallibilismo", "Funcionalismo", "Habitus", 
        "Interacción", "Interdependencia", "Interpretación", "Intuición", "Investigación", "Justificación", 
        "Lenguaje", "Mediación", "Método científico", "Naturaleza", "Objeción", "Objetividad", "Observación", 
        "Operacionalismo", "Paradigma", "Pensamiento", "Percepción", "Pluralismo", "Práctica", "Pragmática", 
        "Pragmatismo", "Proceso", "Producto", "Psicología", "Quehacer", "Racionalidad", "Razón práctica", 
        "Realismo", "Relativismo", "Relación", "Representación", "Resiliencia", "Resolución", "Responsabilidad", 
        "Sensibilidad", "Significado", "Signo", "Simbiosis", "Síntesis", "Sujeto conocedor", "Suposición", 
        "Teoría", "Transformación", "Utopía", "Valor", "Verdad", "Vitalidad", "Voluntarismo"
    ])

    # Pragmatist authors
    autores_pragmatistas = [
        "Charles Sanders Peirce", "William James", "John Dewey", "Richard Rorty", "Hilary Putnam", 
        "George Herbert Mead", "Sidney Hook", "Clarence Irving Lewis", "Josiah Royce", "Jane Addams"
    ]

    def buscar_informacion(query, autor):
        url = "https://google.serper.dev/search"
        payload = json.dumps({
            "q": f"{query} {autor} Filosofía Pragmatista"
        })
        headers = {
            'X-API-KEY': SERPER_API_KEY,
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()

    def generar_definicion(termino, autor, contexto):
        url = "https://api.together.xyz/inference"
        payload = json.dumps({
            "model": "mistralai/Mixtral-8x7B-Instruct-v0.1",
            "prompt": f"Contexto: {contexto}\n\nTérmino: {termino}\nAutor: {autor}\n\nProporciona una definición del término filosófico '{termino}' según el pensamiento de {autor}, un autor de la Filosofía Pragmatista. La definición debe ser concisa pero informativa, similar a una entrada de diccionario. Si es posible, incluye una referencia a una obra específica de {autor} que trate este concepto.\n\nDefinición:",
            "max_tokens": 2048,
            "temperature": 0,
            "top_p": 0.7,
            "top_k": 50,
            "repetition_penalty": 0,
            "stop": ["Término:"]
        })
        headers = {
            'Authorization': f'Bearer {TOGETHER_API_KEY}',
            'Content-Type': 'application/json'
        }
        response = requests.post(url, headers=headers, data=payload)
        return response.json()['output']['choices'][0]['text'].strip()

    def create_docx(termino, definiciones, fuentes):
        doc = Document()
        doc.add_heading('Diccionario Filosófico - Pragmatismo', 0)

        doc.add_heading('Término', level=1)
        doc.add_paragraph(termino)

        for autor, definicion in definiciones.items():
            doc.add_heading(f'Definición según {autor}', level=2)
            doc.add_paragraph(definicion)

        doc.add_heading('Fuentes', level=1)
        for fuente in fuentes:
            doc.add_paragraph(fuente, style='List Bullet')

        doc.add_paragraph('\nNota: Este documento fue generado por un asistente de IA. Verifica la información con fuentes académicas para un análisis más profundo.')

        return doc

    st.write("Elige un término filosófico de la lista o propón tu propio término:")

    opcion = st.radio("", ["Elegir de la lista", "Proponer mi propio término"])

    if opcion == "Elegir de la lista":
        termino = st.selectbox("Selecciona un término:", terminos_filosoficos)
    else:
        termino = st.text_input("Ingresa tu propio término filosófico:")

    st.write("Selecciona uno o más autores pragmatistas (máximo 5):")
    autores_seleccionados = st.multiselect("Autores", autores_pragmatistas)

    if len(autores_seleccionados) > 5:
        st.warning("Has seleccionado más de 5 autores. Por favor, selecciona un máximo de 5.")
    else:
        if st.button("Obtener definición"):
            if termino and autores_seleccionados:
                with st.spinner("Buscando información y generando definiciones..."):
                    definiciones, todas_fuentes = {}, []

                    for autor in autores_seleccionados:
                        # Buscar información relevante
                        resultados_busqueda = buscar_informacion(termino, autor)
                        contexto = "\n".join([item["snippet"] for item in resultados_busqueda.get("organic", [])])
                        fuentes = [item["link"] for item in resultados_busqueda.get("organic", [])]

                        # Generar definición
                        definicion = generar_definicion(termino, autor, contexto)

                        definiciones[autor] = definicion
                        todas_fuentes.extend(fuentes)

                    # Mostrar las definiciones
                    st.subheader(f"Definiciones para el término: {termino}")
                    for autor, definicion in definiciones.items():
                        st.markdown(f"**{autor}:** {definicion}")

                    # Botón para descargar el documento
                    doc = create_docx(termino, definiciones, todas_fuentes)
                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    st.download_button(
                        label="Descargar definición en DOCX",
                        data=buffer,
                        file_name=f"Definicion_{termino.replace(' ', '_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.warning("Por favor, selecciona un término y al menos un autor.")
